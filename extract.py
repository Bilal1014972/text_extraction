import pytesseract
import platform
import shutil
import os
import json
import httpx
from fastapi import APIRouter, UploadFile, File, HTTPException
import fitz  # pymupdf
import pdfplumber
from docx import Document
from PIL import Image
import io
import tempfile
from extraction_prompt import EXTRACTION_SYSTEM_PROMPT, EXTRACTION_USER_PROMPT_TEMPLATE
from pillow_heif import register_heif_opener
from dotenv import load_dotenv
from nutrient_normalization import normalize_nutrition_to_100g

# Enable HEIC/HEIF support (iPhone photos)
register_heif_opener()

load_dotenv()

if platform.system() == "Windows":
    pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
else:
    tesseract_path = shutil.which("tesseract")
    if tesseract_path:
        pytesseract.pytesseract.tesseract_cmd = tesseract_path

router = APIRouter()

ALLOWED_TYPES = {
    "application/pdf",
    "image/png", "image/jpeg", "image/webp", "image/tiff",
    "image/heic", "image/heif",  # iPhone native formats
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
MAX_SIZE = 20 * 1024 * 1024

# HEIC/HEIF brand codes per ISO 23008-12
_HEIC_BRANDS = {b"heic", b"heix", b"hevc", b"hevx", b"heim", b"heis", b"hevm", b"hevs"}
_HEIF_BRANDS = {b"mif1", b"msf1"}

OLLAMA_URL = os.getenv("OLLAMA_URL")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL")
OLLAMA_API_KEY = os.getenv("OLLAMA_API_KEY")


def sniff_mime(data: bytes, declared: str) -> str:
    """
    Detect the true MIME type from magic bytes.
    Falls back to declared Content-Type when no signature matches.
    Prevents rejecting valid files sent as application/octet-stream
    (common with HEIC files in Postman, curl, and browsers).
    """
    if data[:4] == b"%PDF":
        return "application/pdf"
    if data[:4] == b"PK\x03\x04":
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if data[:3] == b"\xff\xd8\xff":
        return "image/jpeg"
    if data[:8] == b"\x89PNG\r\n\x1a\n":
        return "image/png"
    if data[:4] == b"RIFF" and data[8:12] == b"WEBP":
        return "image/webp"
    if data[:4] in (b"II*\x00", b"MM\x00*"):
        return "image/tiff"
    # HEIC/HEIF: ftyp box at byte offset 4
    if data[4:8] == b"ftyp":
        brand = data[8:12]
        if brand in _HEIC_BRANDS:
            return "image/heic"
        if brand in _HEIF_BRANDS:
            return "image/heif"
    return declared


@router.post("/extract")
async def analyze(file: UploadFile = File(...)):
    data = await file.read()

    # Sniff real MIME — must happen before size check and routing
    content_type = sniff_mime(data, file.content_type or "")

    if content_type not in ALLOWED_TYPES:
        raise HTTPException(400, f"Unsupported type: {content_type}")

    if len(data) > MAX_SIZE:
        raise HTTPException(413, "File too large (max 20MB)")

    # Step 1: Extract text — use sniffed content_type, NOT file.content_type
    try:
        if content_type == "application/pdf":
            extraction = extract_from_pdf(data)
        elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            extraction = extract_from_docx(data)
        else:
            extraction = {"pages": [{"page": 1, "text": extract_from_image(data)}]}
    except Exception as e:
        raise HTTPException(500, f"Extraction failed: {str(e)}")

    all_text = "\n\n".join(
        page["text"] for page in extraction["pages"] if page.get("text")
    )

    if not all_text.strip():
        raise HTTPException(422, "No text could be extracted from the document")

    # Step 2: LLM analysis
    try:
        structured_data = await call_ollama(all_text)
    except json.JSONDecodeError:
        raise HTTPException(502, "LLM returned invalid JSON")
    except Exception as e:
        raise HTTPException(502, f"LLM analysis failed: {str(e)}")
    
    # Step 3: Normalize nutrition to 100g  ← ADD THIS
    structured_data = normalize_nutrition_to_100g(structured_data)

    return {
        "filename": file.filename,
        **structured_data,
    }


async def call_ollama(extracted_text: str) -> dict:
    headers = {"Content-Type": "application/json"}
    if OLLAMA_API_KEY:
        headers["Authorization"] = f"Bearer {OLLAMA_API_KEY}"

    payload = {
        "model": OLLAMA_MODEL,
        "messages": [
            {"role": "system", "content": EXTRACTION_SYSTEM_PROMPT},
            {
                "role": "user",
                "content": EXTRACTION_USER_PROMPT_TEMPLATE.format(
                    extracted_text=extracted_text
                ),
            },
        ],
        "stream": False,
        "options": {
            "temperature": 0.4,
            "num_ctx": 16384,
        },
    }

    async with httpx.AsyncClient(timeout=120.0) as client:
        response = await client.post(
            f"{OLLAMA_URL}/api/chat",
            json=payload,
            headers=headers,
        )
        response.raise_for_status()

    result = response.json()
    content = result["message"]["content"]

    content = content.strip()
    if content.startswith("```json"):
        content = content[7:]
    if content.startswith("```"):
        content = content[3:]
    if content.endswith("```"):
        content = content[:-3]
    content = content.strip()

    return json.loads(content)


def extract_from_pdf(data: bytes) -> dict:
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(data)
        tmp_path = tmp.name

    all_pages = []
    try:
        with pdfplumber.open(tmp_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text() or ""
                if not text.strip():
                    text = ocr_page(data, page_num - 1)
                all_pages.append({"page": page_num, "text": text.strip()})
    finally:
        os.unlink(tmp_path)

    return {"pages": all_pages}


def extract_from_docx(data: bytes) -> dict:
    doc = Document(io.BytesIO(data))
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_rows.append(" | ".join(cells))
        parts.append("\n".join(table_rows))

    return {"pages": [{"page": 1, "text": "\n".join(parts)}]}


def ocr_page(pdf_data: bytes, page_index: int) -> str:
    doc = fitz.open(stream=pdf_data, filetype="pdf")
    page = doc[page_index]
    pix = page.get_pixmap(dpi=300)
    img = Image.open(io.BytesIO(pix.tobytes("png")))
    text = pytesseract.image_to_string(img)
    doc.close()
    return text.strip()


def extract_from_image(data: bytes) -> str:
    img = Image.open(io.BytesIO(data))
    # Pytesseract checks image.format against a whitelist (PNG, JPEG, TIFF...).
    # HEIF is not on that list even when mode is already RGB.
    # Fix: re-encode through an in-memory PNG buffer to reset the format tag.
    buf = io.BytesIO()
    img.convert("RGB").save(buf, format="PNG")
    buf.seek(0)
    img = Image.open(buf)
    return pytesseract.image_to_string(img)