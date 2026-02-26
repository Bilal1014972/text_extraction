import pytesseract
import os
import json
import httpx
from fastapi import FastAPI, UploadFile, File, HTTPException
import fitz  # pymupdf
import pdfplumber
from docx import Document
from PIL import Image
import io
import tempfile
from extraction_prompt import EXTRACTION_SYSTEM_PROMPT, EXTRACTION_USER_PROMPT_TEMPLATE

app = FastAPI()

ALLOWED_TYPES = {
    "application/pdf",
    "image/png", "image/jpeg", "image/webp", "image/tiff",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
MAX_SIZE = 20 * 1024 * 1024

OLLAMA_URL = os.getenv("OLLAMA_URL")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL")
OLLAMA_API_KEY = os.getenv("OLLAMA_API_KEY")


@app.post("/extract")
async def analyze(file: UploadFile = File(...)):
    if file.content_type not in ALLOWED_TYPES:
        raise HTTPException(400, f"Unsupported type: {file.content_type}")

    data = await file.read()
    if len(data) > MAX_SIZE:
        raise HTTPException(413, "File too large (max 20MB)")

    # Step 1: Extract text
    try:
        if file.content_type == "application/pdf":
            extraction = extract_from_pdf(data)
        elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            extraction = extract_from_docx(data)
        else:
            extraction = {"pages": [{"page": 1, "text": extract_from_image(data)}]}
    except Exception as e:
        raise HTTPException(500, f"Extraction failed: {str(e)}")

    all_text = "\n\n".join(
        page["text"] for page in extraction["pages"] if page.get("text")
    )

    if not all_text.strip():
        raise HTTPException(422, "No text could be extracted from the document")

    # Step 2: LLM analysis
    try:
        structured_data = await call_ollama(all_text)
    except json.JSONDecodeError:
        raise HTTPException(502, "LLM returned invalid JSON")
    except Exception as e:
        raise HTTPException(502, f"LLM analysis failed: {str(e)}")

    return {
        "filename": file.filename,
        "structured_data": structured_data,
    }


async def call_ollama(extracted_text: str) -> dict:
    headers = {"Content-Type": "application/json"}
    if OLLAMA_API_KEY:
        headers["Authorization"] = f"Bearer {OLLAMA_API_KEY}"

    payload = {
        "model": OLLAMA_MODEL,
        "messages": [
            {"role": "system", "content": EXTRACTION_SYSTEM_PROMPT},
            {
                "role": "user",
                "content": EXTRACTION_USER_PROMPT_TEMPLATE.format(
                    extracted_text=extracted_text
                ),
            },
        ],
        "stream": False,
        "options": {
            "temperature": 0.4,
            "num_ctx": 16384,
        },
    }

    async with httpx.AsyncClient(timeout=120.0) as client:
        response = await client.post(
            f"{OLLAMA_URL}/api/chat",
            json=payload,
            headers=headers,
        )
        response.raise_for_status()

    result = response.json()
    content = result["message"]["content"]

    content = content.strip()
    if content.startswith("```json"):
        content = content[7:]
    if content.startswith("```"):
        content = content[3:]
    if content.endswith("```"):
        content = content[:-3]
    content = content.strip()

    return json.loads(content)


def extract_from_pdf(data: bytes) -> dict:
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(data)
        tmp_path = tmp.name

    all_pages = []
    try:
        with pdfplumber.open(tmp_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text() or ""
                if not text.strip():
                    text = ocr_page(data, page_num - 1)
                all_pages.append({"page": page_num, "text": text.strip()})
    finally:
        os.unlink(tmp_path)

    return {"pages": all_pages}


def extract_from_docx(data: bytes) -> dict:
    doc = Document(io.BytesIO(data))
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_rows.append(" | ".join(cells))
        parts.append("\n".join(table_rows))

    return {"pages": [{"page": 1, "text": "\n".join(parts)}]}


def ocr_page(pdf_data: bytes, page_index: int) -> str:
    doc = fitz.open(stream=pdf_data, filetype="pdf")
    page = doc[page_index]
    pix = page.get_pixmap(dpi=300)
    img = Image.open(io.BytesIO(pix.tobytes("png")))
    text = pytesseract.image_to_string(img)
    doc.close()
    return text.strip()


def extract_from_image(data: bytes) -> str:
    img = Image.open(io.BytesIO(data))
    return pytesseract.image_to_string(img)